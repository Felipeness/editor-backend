from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.text.run import Run
from ..utils import add_hyperlink

def _text_runs(node, run_style):
    out = []
    if isinstance(node, NavigableString):
        text = str(node)
        if text:
            out.append((text, run_style, None))
        return out

    if not isinstance(node, Tag):
        return out

    tag = node.name.lower()
    next_style = dict(run_style)
    href = None

    if tag in ("b", "strong"):
        next_style["bold"] = True
    if tag in ("i", "em"):
        next_style["italic"] = True
    if tag == "a":
        href = node.get("href")

    if tag == "a":
        # compacta como um único run com href
        text = node.get_text("", strip=False)
        out.append((text, next_style, href))
        return out

    for child in node.children:
        out.extend(_text_runs(child, next_style))
    return out

def _apply_runs(paragraph, node, base_size_pt: int | None = None):
    runs = _text_runs(node, {"bold": False, "italic": False, "size": base_size_pt})
    if not runs:
        paragraph.add_run("")
        return

    for text, style, href in runs:
        if href:
            add_hyperlink(paragraph, href, text or href)
            continue
        r: Run = paragraph.add_run(text)
        if style.get("bold"):   r.bold = True
        if style.get("italic"): r.italic = True
        if style.get("size"):   r.font.size = Pt(style["size"])

def _alignment_from_css(tag: Tag):
    s = (tag.get("style") or "").lower()
    if "text-align:center" in s:  return WD_ALIGN_PARAGRAPH.CENTER
    if "text-align:right" in s:   return WD_ALIGN_PARAGRAPH.RIGHT
    if "text-align:justify" in s: return WD_ALIGN_PARAGRAPH.JUSTIFY
    return None

def _list_level(li: Tag) -> int:
    level = 0; p = li.parent
    while isinstance(p, Tag):
        if p.name and p.name.lower() in ("ul", "ol"):
            level += 1
        p = p.parent
    return max(0, level - 1)

def _is_li_heading(li: Tag) -> int | None:
    dh = li.get("data-heading")
    if dh in ("1","2","3"):
        return int(dh)
    return None

def html_to_docx(html: str, title: str, author: str) -> bytes:
    soup = BeautifulSoup(html or "", "lxml")
    doc = Document()

    section = doc.sections[0]
    section.page_width  = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin    = Inches(0.79)
    section.bottom_margin = Inches(0.79)
    section.left_margin   = Inches(0.79)
    section.right_margin  = Inches(0.79)

    doc.core_properties.title = title or "Untitled"
    doc.core_properties.author = author or "Anonymous"

    def handle_block(tag: Tag):
        name = tag.name.lower()
        align = _alignment_from_css(tag)

        if name in ("h1","h2","h3"):
            p = doc.add_paragraph()
            p.style = {"h1":"Heading 1","h2":"Heading 2","h3":"Heading 3"}[name]
            base = {"h1":24, "h2":18, "h3":14}[name]
            _apply_runs(p, tag, base_size_pt=base)
            if align: p.alignment = align
            return

        if name == "blockquote":
            for child in tag.children:
                if isinstance(child, NavigableString):
                    text = str(child).strip()
                    if not text: continue
                    p = doc.add_paragraph(text)
                    p.paragraph_format.left_indent = Inches(0.5)
                elif isinstance(child, Tag):
                    p = doc.add_paragraph()
                    _apply_runs(p, child)
                    p.paragraph_format.left_indent = Inches(0.5)
            return

        if name in ("ul","ol"):
            ordered = name == "ol"
            for li in tag.select(":scope > li"):
                lvl = _list_level(li)
                heading_level = _is_li_heading(li)

                p = doc.add_paragraph()
                p.style = "List Number" if ordered else "List Bullet"
                if lvl > 0:
                    p.paragraph_format.left_indent = Inches(0.3 * (lvl + 1))

                base = None
                if heading_level == 1: base = 24
                elif heading_level == 2: base = 18
                elif heading_level == 3: base = 14

                li_text = li.find("span", {"class":"li-text"})
                node = li_text or li
                _apply_runs(p, node, base_size_pt=base)

                # listas aninhadas diretas
                for nested in li.find_all(["ul","ol"], recursive=False):
                    handle_block(nested)
            return

        if name in ("p","div"):
            p = doc.add_paragraph()
            _apply_runs(p, tag)
            if align: p.alignment = align
            return

        p = doc.add_paragraph()
        _apply_runs(p, tag)
        if align: p.alignment = align

    body = soup.body or soup
    blocks = [c for c in body.children if isinstance(c, (Tag, NavigableString))]
    if not any(isinstance(c, Tag) for c in blocks):
        doc.add_paragraph("")

    for node in blocks:
        if isinstance(node, NavigableString):
            text = str(node).strip()
            if text: doc.add_paragraph(text)
            continue
        handle_block(node)

    from io import BytesIO
    buf = BytesIO(); doc.save(buf)
    return buf.getvalue()
